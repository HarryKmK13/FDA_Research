"""Build the standalone DOCX writing sample from the Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "writing_sample" / "Kyaw_Min_Khant_FDA_Citizen_Petitions_Research_Paper.md"
OUTPUT = ROOT / "reports" / "writing_sample" / "Kyaw_Min_Khant_FDA_Citizen_Petitions_Research_Paper.docx"
FIGURE_DIR = ROOT / "figures"
DOCS_SKILL = Path(
    "/Users/harryhuang/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.430.10722/skills/documents/scripts"
)
sys.path.append(str(DOCS_SKILL))
from table_geometry import apply_table_geometry  # noqa: E402


ACCENT = RGBColor(47, 111, 115)
MUTED = RGBColor(92, 101, 107)
LIGHT_FILL = "E8F1F1"
BORDER = "B8C7C9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_bottom_border(paragraph, color: str = BORDER) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def setup_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(10)

    for name, size, before, after in [
        ("Heading 1", 14, 12, 5),
        ("Heading 2", 12, 9, 4),
        ("Heading 3", 11, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    quote = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    quote.font.name = "Arial"
    quote.font.size = Pt(10)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor(64, 75, 79)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.25)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(8)

    caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)


def add_run_with_inline_formatting(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        run = paragraph.add_run(token[2:-2] if token.startswith("**") else token[1:-1])
        if token.startswith("**"):
            run.bold = True
        else:
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_metadata_table(document: Document) -> None:
    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Research area", "FDA citizen petitions, public regulatory records, and agency response timing"),
        ("Supervisor", "Allison Schmitt, aaschmit@uoregon.edu"),
        ("Dataset", "2011 FDA citizen petition and response PDFs converted into validated structured tables"),
        ("Core methods", "PDF/OCR extraction, local LLM field extraction, docket matching, response categorization, visualization"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                paragraph.paragraph_format.space_after = Pt(0)
        row.cells[0].paragraphs[0].runs[0].bold = True
    apply_table_geometry(table, [2100, 7260], table_width_dxa=9360, indent_dxa=0)
    document.add_paragraph()


def add_key_metrics_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Metric", "Value", "Interpretation"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        set_cell_shading(table.rows[0].cells[idx], LIGHT_FILL)
        table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
    rows = [
        ("Raw document metadata records", "652", "Public docket materials identified before filtering."),
        ("Validated petition rows", "129", "Structured petition records available for analysis."),
        ("Validated response rows", "120", "Structured response records available for outcome coding."),
        ("Valid response-time pairs", "64", "Nonnegative petition-response row matches after date QA."),
        ("Median response time", "182.5 days", "Typical valid pair is near the 180-day general response window."),
        ("Most common response category", "Interim", "Many records indicate continued review rather than final resolution."),
    ]
    for metric, value, interpretation in rows:
        cells = table.add_row().cells
        cells[0].text = metric
        cells[1].text = value
        cells[2].text = interpretation
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.8)
            if idx == 1:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_table_geometry(table, [2800, 1350, 5210], table_width_dxa=9360, indent_dxa=0)
    document.add_paragraph()


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))
    caption_paragraph = document.add_paragraph(caption, style="Figure Caption")
    caption_paragraph.paragraph_format.keep_together = True


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    setup_styles(document)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Kyaw Min (Harry) Khant | FDA Citizen Petition Research Paper"
    header_para.style = document.styles["Normal"]
    header_para.runs[0].font.size = Pt(8.5)
    header_para.runs[0].font.color.rgb = MUTED
    add_bottom_border(header_para)
    add_page_number(section.footer.paragraphs[0])

    text = SOURCE.read_text(encoding="utf-8").splitlines()

    title = text[0].lstrip("# ").strip()
    title_para = document.add_paragraph(title, style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(
        "Kyaw Min (Harry) Khant | Supervised by Allison Schmitt (aaschmit@uoregon.edu)",
        style="Subtitle",
    )
    add_metadata_table(document)

    skip_until = 1
    while skip_until < len(text) and not text[skip_until].startswith("## Abstract"):
        skip_until += 1

    in_references = False
    pending_number = 1
    paragraph_buffer: list[str] = []

    def flush_buffer() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        add_run_with_inline_formatting(paragraph, " ".join(paragraph_buffer).strip())
        paragraph_buffer = []

    for raw_line in text[skip_until:]:
        line = raw_line.strip()
        if not line:
            flush_buffer()
            continue
        if line.startswith("!["):
            flush_buffer()
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                caption = match.group(1)
                relative = match.group(2)
                image_path = (SOURCE.parent / relative).resolve()
                add_figure(document, image_path, caption)
            continue
        if line.startswith("## "):
            flush_buffer()
            heading = line[3:].strip()
            in_references = heading == "References"
            document.add_paragraph(heading, style="Heading 1")
            if heading == "Findings":
                add_key_metrics_table(document)
            continue
        if line.startswith("### "):
            flush_buffer()
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            continue
        if re.match(r"\d+\.\s+", line):
            flush_buffer()
            content = re.sub(r"^\d+\.\s+", "", line)
            paragraph = document.add_paragraph(style="List Number")
            add_run_with_inline_formatting(paragraph, content)
            pending_number += 1
            continue
        if line.startswith("- "):
            flush_buffer()
            paragraph = document.add_paragraph(style="List Bullet")
            add_run_with_inline_formatting(paragraph, line[2:])
            continue
        if in_references:
            flush_buffer()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            paragraph.paragraph_format.space_after = Pt(5)
            add_run_with_inline_formatting(paragraph, line)
            continue
        paragraph_buffer.append(line)

    flush_buffer()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
